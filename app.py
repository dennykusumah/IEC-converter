"""
engine14.py – DOCX Post-Processor Engine (FIXED TABLE LAYOUT)
Memperbaiki file DOCX keluaran dari engine13:
1. Mengubah semua gambar menjadi wrap "In line with text"
2. Format keterangan gambar: Arial 10pt, single, 0pt spacing, UNBOLD
3. Kecuali tulisan "Key" atau "NOTE/Note to entry" yang tetap BOLD
4. Memperbaiki tabel yang melebihi ukuran kertas (Autoscale Grid & Shrink Font)
"""

import re
import copy
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import logging

logger = logging.getLogger(__name__)

class DOCXPostProcessorEngine:

    def process(self, docx_path: str) -> tuple:
        try:
            doc = Document(docx_path)
            
            logger.info("Tahap 4.1: Mengubah gambar ke inline...")
            self._set_images_inline(doc)
            
            logger.info("Tahap 4.2: Memformat caption gambar...")
            self._process_captions(doc)
            
            logger.info("Tahap 4.3: Memperbaiki tabel overflow...")
            self._fix_table_overflow(doc)
            
            doc.save(docx_path)
            return True, docx_path, {"status": "Post-processing berhasil"}
            
        except Exception as e:
            logger.error(f"Gagal post-processing: {e}")
            return False, str(e), {}

    # ══════════════════════════════════════════════════════════════════════
    # 1. GAMBAR: Ubah Anchor (Floating) menjadi Inline
    # ══════════════════════════════════════════════════════════════════════
    def _set_images_inline(self, doc):
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        }
        
        # Proses paragraf biasa
        for para in doc.paragraphs:
            for anchor in para._element.xpath('//w:drawing/wp:anchor', namespaces=namespaces):
                self._replace_anchor_with_inline(anchor, namespaces)
                
        # Proses gambar yang ada di dalam tabel
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for anchor in para._element.xpath('//w:drawing/wp:anchor', namespaces=namespaces):
                            self._replace_anchor_with_inline(anchor, namespaces)

    def _replace_anchor_with_inline(self, anchor, ns):
        extent = anchor.find('wp:extent', ns)
        effectExtent = anchor.find('wp:effectExtent', ns)
        docPr = anchor.find('wp:docPr', ns)
        cNvGraphicFramePr = anchor.find('wp:cNvGraphicFramePr', ns)
        graphic = anchor.find('a:graphic', ns)
        
        if extent is None or docPr is None or graphic is None:
            return
            
        inline = OxmlElement('wp:inline')
        inline.set('distT', '0')
        inline.set('distB', '0')
        inline.set('distL', '0')
        inline.set('distR', '0')
        
        inline.append(copy.deepcopy(extent))
        if effectExtent is not None:
            inline.append(copy.deepcopy(effectExtent))
        inline.append(copy.deepcopy(docPr))
        if cNvGraphicFramePr is not None:
            inline.append(copy.deepcopy(cNvGraphicFramePr))
        inline.append(copy.deepcopy(graphic))
        
        parent = anchor.getparent()
        parent.remove(anchor)
        parent.append(inline)

    # ══════════════════════════════════════════════════════════════════════
    # 2. CAPTION: Deteksi dan Format Blok Gambar (AMAN UNTUK GAMBAR INLINE)
    # ══════════════════════════════════════════════════════════════════════
    def _process_captions(self, doc):
        paras = doc.paragraphs
        i = 0
        while i < len(paras):
            text = paras[i].text.strip()
            
            if re.match(r'^Figure \d+', text, re.IGNORECASE):
                self._format_caption_text(paras[i], text)
                i += 1
                
                while i < len(paras):
                    next_text = paras[i].text.strip()
                    if re.match(r'^Key$', next_text, re.IGNORECASE):
                        self._format_caption_text(paras[i], next_text)
                        i += 1
                        while i < len(paras):
                            item_text = paras[i].text.strip()
                            if re.match(r'^\d+\s+.+', item_text):
                                self._format_caption_text(paras[i], item_text)
                                i += 1
                            elif not item_text:
                                i += 1
                            else:
                                break
                    elif not next_text:
                        i += 1
                    else:
                        break
                        
            elif re.match(r'^(NOTE|Note\s*\d*\s*to\s*entry\s*:)', text, re.IGNORECASE):
                self._format_caption_text(paras[i], text)
                i += 1
            else:
                i += 1

    def _format_caption_text(self, para, text):
        if not text:
            return
            
        p = para._element
        pPr = p.get_or_add_pPr()
        
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        spacing.set(qn('w:line'), '240')
        spacing.set(qn('w:lineRule'), 'auto')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        
        # PENTING: Hanya hapus run TEKS, jangan hapus run GAMBAR (w:drawing)
        runs_to_remove = []
        for r in p.findall(qn('w:r')):
            if r.find(qn('w:drawing')) is None and r.find(qn('w:pict')) is None:
                runs_to_remove.append(r)
        for r in runs_to_remove:
            p.remove(r)
            
        def add_run(txt, is_bold):
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            
            if is_bold:
                b = OxmlElement('w:b')
                rPr.append(b)
                bCs = OxmlElement('w:bCs')
                rPr.append(bCs)
                
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Arial')
            rFonts.set(qn('w:hAnsi'), 'Arial')
            rFonts.set(qn('w:cs'), 'Arial')
            rPr.append(rFonts)
            
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '20')
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '20')
            rPr.append(szCs)
            
            r.append(rPr)
            t = OxmlElement('w:t')
            t.text = txt
            t.set(qn('xml:space'), 'preserve')
            r.append(t)
            p.append(r)

        m_key = re.match(r'^(Key)(.*)', text, re.IGNORECASE)
        if m_key:
            add_run(m_key.group(1), True)
            if m_key.group(2):
                add_run(m_key.group(2), False)
            return

        m_note = re.match(r'^(NOTE|Note\s*\d*\s*to\s*entry\s*:)(.*)', text, re.IGNORECASE)
        if m_note:
            add_run(m_note.group(1), True)
            if m_note.group(2):
                add_run(m_note.group(2), False)
            return

        add_run(text, False)

    # ══════════════════════════════════════════════════════════════════════
    # 3. TABEL: Perbaiki Tabel yang Melebihi Ukuran Kertas (ROBUST MERGE)
    # ══════════════════════════════════════════════════════════════════════
    def _fix_table_overflow(self, doc):
        PAGE_WIDTH_TWIPS = 9072  # A4 (21cm) - Margin (2.5cm x 2) = 16cm
        
        for table in doc.tables:
            tbl = table._tbl
            tblPr = tbl.get_or_add_tblPr()
            
            # 1. Normalisasi lebar Grid Kolom dasar
            tblGrid = tbl.find(qn('w:tblGrid'))
            if tblGrid is not None:
                gridCols = tblGrid.findall(qn('w:gridCol'))
                if gridCols:
                    total_width = sum(int(col.get(qn('w:w'), 0)) for col in gridCols)
                    # Jika total lebar melebihi kertas, pangkatkan proporsional
                    if total_width > PAGE_WIDTH_TWIPS and total_width > 0:
                        scale = PAGE_WIDTH_TWIPS / total_width
                        for col in gridCols:
                            new_w = int(int(col.get(qn('w:w'), 0)) * scale)
                            col.set(qn('w:w'), str(new_w))
            
            # 2. Set Layout ke FIXED (Mutlak mematuhi Grid agar Merge Cell tidak rusak)
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is None:
                tblW = OxmlElement('w:tblW')
                tblPr.append(tblW)
            tblW.set(qn('w:w'), str(PAGE_WIDTH_TWIPS))
            tblW.set(qn('w:type'), 'dxa')
            
            tblLayout = tblPr.find(qn('w:tblLayout'))
            if tblLayout is None:
                tblLayout = OxmlElement('w:tblLayout')
                tblPr.append(tblLayout)
            tblLayout.set(qn('w:type'), 'fixed') 
            
            # 3. Proses Setiap Sel
            for row in table.rows:
                for cell in row.cells:
                    tcPr = cell._element.get_or_add_tcPr()
                    
                    # HAPUS lebar eksplisit sel (biar Grid yang ngatur, mencegah bug merge cell)
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is not None:
                        tcPr.remove(tcW)
                    
                    # Set Margin Sel menjadi super ketat (40 twips ≈ 0.07cm) agar kolom muat banyak
                    tcMar = tcPr.find(qn('w:tcMar'))
                    if tcMar is not None:
                        tcPr.remove(tcMar)
                    new_tcMar = OxmlElement('w:tcMar')
                    for pos in ['top', 'left', 'bottom', 'right']:
                        margin = OxmlElement(f'w:{pos}')
                        margin.set(qn('w:w'), '40')
                        margin.set(qn('w:type'), 'dxa')
                        new_tcMar.append(margin)
                    tcPr.append(new_tcMar)
                    
                    # Pastikan teks bisa wrap (tidak terpotong di luar sel)
                    noWrap = tcPr.find(qn('w:noWrap'))
                    if noWrap is not None:
                        tcPr.remove(noWrap)
                    
                    # 4. Format Paragraf & Font di dalam sel
                    for para in cell.paragraphs:
                        pPr = para._element.get_or_add_pPr()
                        
                        # Hapus indentasi asal dari PDF yang bikin text makin terdorong dalam
                        ind = pPr.find(qn('w:ind'))
                        if ind is not None:
                            pPr.remove(ind)
                            
                        # Set spacing jadi rapat
                        spacing = pPr.find(qn('w:spacing'))
                        if spacing is None:
                            spacing = OxmlElement('w:spacing')
                            pPr.append(spacing)
                        spacing.set(qn('w:before'), '0')
                        spacing.set(qn('w:after'), '0')
                        spacing.set(qn('w:line'), '200') # 10pt line spacing (rapat)
                        spacing.set(qn('w:lineRule'), 'exact')
                        
                        # Paksa font jadi 8pt (standar tabel IEC yang rapi)
                        for run in para.runs:
                            if run.font.size is None or run.font.size > Pt(8):
                                run.font.size = Pt(8)
