"""
engine14.py – DOCX Post-Processor Engine (ADVANCED AUTO-FIT TABLE)
Memperbaiki file DOCX keluaran dari engine13:
1. Mengubah semua gambar menjadi wrap "In line with text"
2. Format keterangan gambar: Arial 10pt, single, 0pt spacing, UNBOLD
3. Kecuali tulisan "Key" atau "NOTE/Note to entry" yang tetap BOLD
4. Memperbaiki tabel melebihi ukuran kertas secara CERDAS:
   - Menggunakan logika "Auto-fit to Window" (100% margin)
   - Menghapus lebar kolom kaku dari PDF agar Word yang mengatur
   - Mengecilkan font secara otomatis berdasarkan kepadatan kolom
"""

import re
import copy
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import logging

logger = logging.getLogger(__name__)

class DOCXPostProcessorEngine:

    def process(self, docx_path: str) -> tuple:
        try:
            doc = Document(docx_path)
            
            logger.info("Tahap 4.1: Mengubah gambar ke inline...")
            self._set_images_inline(doc)
            
            logger.info("Tahap 4.2: Memformat caption gambar...")
            self._process_captions(doc)
            
            logger.info("Tahap 4.3: Memperbaiki tabel overflow (Auto-Fit)...")
            self._fix_table_overflow(doc)
            
            doc.save(docx_path)
            return True, docx_path, {"status": "Post-processing berhasil"}
            
        except Exception as e:
            logger.error(f"Gagal post-processing: {e}")
            return False, str(e), {}

    # ══════════════════════════════════════════════════════════════════════
    # 1. GAMBAR: Ubah Anchor (Floating) menjadi Inline
    # ══════════════════════════════════════════════════════════════════════
    def _set_images_inline(self, doc):
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        }
        
        for para in doc.paragraphs:
            for anchor in para._element.xpath('//w:drawing/wp:anchor', namespaces=namespaces):
                self._replace_anchor_with_inline(anchor, namespaces)
                
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for anchor in para._element.xpath('//w:drawing/wp:anchor', namespaces=namespaces):
                            self._replace_anchor_with_inline(anchor, namespaces)

    def _replace_anchor_with_inline(self, anchor, ns):
        extent = anchor.find('wp:extent', ns)
        effectExtent = anchor.find('wp:effectExtent', ns)
        docPr = anchor.find('wp:docPr', ns)
        cNvGraphicFramePr = anchor.find('wp:cNvGraphicFramePr', ns)
        graphic = anchor.find('a:graphic', ns)
        
        if extent is None or docPr is None or graphic is None:
            return
            
        inline = OxmlElement('wp:inline')
        inline.set('distT', '0')
        inline.set('distB', '0')
        inline.set('distL', '0')
        inline.set('distR', '0')
        
        inline.append(copy.deepcopy(extent))
        if effectExtent is not None:
            inline.append(copy.deepcopy(effectExtent))
        inline.append(copy.deepcopy(docPr))
        if cNvGraphicFramePr is not None:
            inline.append(copy.deepcopy(cNvGraphicFramePr))
        inline.append(copy.deepcopy(graphic))
        
        parent = anchor.getparent()
        parent.remove(anchor)
        parent.append(inline)

    # ══════════════════════════════════════════════════════════════════════
    # 2. CAPTION: Deteksi dan Format Blok Gambar
    # ══════════════════════════════════════════════════════════════════════
    def _process_captions(self, doc):
        paras = doc.paragraphs
        i = 0
        while i < len(paras):
            text = paras[i].text.strip()
            
            if re.match(r'^Figure \d+', text, re.IGNORECASE):
                self._format_caption_text(paras[i], text)
                i += 1
                
                while i < len(paras):
                    next_text = paras[i].text.strip()
                    if re.match(r'^Key$', next_text, re.IGNORECASE):
                        self._format_caption_text(paras[i], next_text)
                        i += 1
                        while i < len(paras):
                            item_text = paras[i].text.strip()
                            if re.match(r'^\d+\s+.+', item_text):
                                self._format_caption_text(paras[i], item_text)
                                i += 1
                            elif not item_text:
                                i += 1
                            else:
                                break
                    elif not next_text:
                        i += 1
                    else:
                        break
                        
            elif re.match(r'^(NOTE|Note\s*\d*\s*to\s*entry\s*:)', text, re.IGNORECASE):
                self._format_caption_text(paras[i], text)
                i += 1
            else:
                i += 1

    def _format_caption_text(self, para, text):
        if not text:
            return
            
        p = para._element
        pPr = p.get_or_add_pPr()
        
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        spacing.set(qn('w:line'), '240')
        spacing.set(qn('w:lineRule'), 'auto')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        
        # Hanya hapus run TEKS, amankan run GAMBAR
        runs_to_remove = []
        for r in p.findall(qn('w:r')):
            if r.find(qn('w:drawing')) is None and r.find(qn('w:pict')) is None:
                runs_to_remove.append(r)
        for r in runs_to_remove:
            p.remove(r)
            
        def add_run(txt, is_bold):
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            
            if is_bold:
                b = OxmlElement('w:b')
                rPr.append(b)
                bCs = OxmlElement('w:bCs')
                rPr.append(bCs)
                
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Arial')
            rFonts.set(qn('w:hAnsi'), 'Arial')
            rFonts.set(qn('w:cs'), 'Arial')
            rPr.append(rFonts)
            
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '20')
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '20')
            rPr.append(szCs)
            
            r.append(rPr)
            t = OxmlElement('w:t')
            t.text = txt
            t.set(qn('xml:space'), 'preserve')
            r.append(t)
            p.append(r)

        m_key = re.match(r'^(Key)(.*)', text, re.IGNORECASE)
        if m_key:
            add_run(m_key.group(1), True)
            if m_key.group(2):
                add_run(m_key.group(2), False)
            return

        m_note = re.match(r'^(NOTE|Note\s*\d*\s*to\s*entry\s*:)(.*)', text, re.IGNORECASE)
        if m_note:
            add_run(m_note.group(1), True)
            if m_note.group(2):
                add_run(m_note.group(2), False)
            return

        add_run(text, False)

    # ══════════════════════════════════════════════════════════════════════
    # 3. TABEL: "Auto-Fit to Window" (Menyesuaikan Margin & Font Dinamis)
    # ══════════════════════════════════════════════════════════════════════
    def _fix_table_overflow(self, doc):
        for table in doc.tables:
            tbl = table._tbl
            tblPr = tbl.get_or_add_tblPr()
            
            # 1. HAPUS semua lebar sel eksplisit (w:tcW) agar tidak rebutan ruang
            for row in table.rows:
                for cell in row.cells:
                    tcPr = cell._element.get_or_add_tcPr()
                    
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is not None:
                        tcPr.remove(tcW)
                    
                    # Set Margin Sel jadi super ketat (0.05cm) agar kolom bisa mengerut
                    tcMar = tcPr.find(qn('w:tcMar'))
                    if tcMar is not None:
                        tcPr.remove(tcMar)
                    new_tcMar = OxmlElement('w:tcMar')
                    for pos in ['top', 'left', 'bottom', 'right']:
                        margin = OxmlElement(f'w:{pos}')
                        margin.set(qn('w:w'), '28') # ~0.05cm (sangat rapat)
                        margin.set(qn('w:type'), 'dxa')
                        new_tcMar.append(margin)
                    tcPr.append(new_tcMar)
                    
                    # Pastikan teks bisa wrap (tidak terpotong)
                    noWrap = tcPr.find(qn('w:noWrap'))
                    if noWrap is not None:
                        tcPr.remove(noWrap)

            # 2. SET MODE TABLE: Auto-Fit to Window (5000 = 100% margin kertas)
            # Ini memerintahkan MS Word untuk MENGABAIKAN lebar PDF, dan merasio 
            # ulang semua kolom agar PAS exactly di dalam margin halaman.
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is None:
                tblW = OxmlElement('w:tblW')
                tblPr.append(tblW)
            tblW.set(qn('w:w'), '5000') 
            tblW.set(qn('w:type'), 'pct')
            
            tblLayout = tblPr.find(qn('w:tblLayout'))
            if tblLayout is None:
                tblLayout = OxmlElement('w:tblLayout')
                tblPr.append(tblLayout)
            tblLayout.set(qn('w:type'), 'autofit')
            
            # 3. UKURAN FONT OTOMATIS berdasarkan kepadatan kolom
            tblGrid = tbl.find(qn('w:tblGrid'))
            num_cols = len(tblGrid.findall(qn('w:gridCol'))) if tblGrid is not None else 0
            
            target_font = Pt(10)
            if num_cols > 20:
                target_font = Pt(5)   # Sangat padat
            elif num_cols > 15:
                target_font = Pt(6)   # Padat
            elif num_cols > 10:
                target_font = Pt(7)   # Agak padat
            elif num_cols > 7:
                target_font = Pt(8)   # Normal IEC
                
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        pPr = para._element.get_or_add_pPr()
                        
                        # Hapus indentasi kaku dari PDF
                        ind = pPr.find(qn('w:ind'))
                        if ind is not None:
                            pPr.remove(ind)
                            
                        # Set line spacing jadi rapat
                        spacing = pPr.find(qn('w:spacing'))
                        if spacing is None:
                            spacing = OxmlElement('w:spacing')
                            pPr.append(spacing)
                        spacing.set(qn('w:before'), '0')
                        spacing.set(qn('w:after'), '0')
                        spacing.set(qn('w:line'), '180') # Sangat rapat
                        spacing.set(qn('w:lineRule'), 'exact')
                        
                        # Paksa turunkan ukuran font jika melebihi batas target
                        for run in para.runs:
                            if run.font.size is None or run.font.size > target_font:
                                run.font.size = target_font