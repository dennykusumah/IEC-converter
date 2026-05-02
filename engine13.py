"""
engine13.py – PDF to DOCX Converter (ENHANCED)
Mendukung PDF digital dan PDF scan (OCR via Tesseract).
Dengan post-processing untuk:
  - Semua gambar → inline with text
  - Caption gambar → Arial 10pt, unbold, single, 0pt spacing
  - Kecuali "Key" atau "NOTE" yang tetap bold
"""

import os
import fitz  # PyMuPDF
from pdf2docx import Converter
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import copy
import logging

logger = logging.getLogger(__name__)


class PDFConverterEngine:

    def __init__(self, tesseract_path=None):
        self.tesseract_path = tesseract_path
        if tesseract_path and os.path.exists(tesseract_path):
            try:
                import pytesseract
                pytesseract.pytesseract.tesseract_cmd = tesseract_path
            except ImportError:
                print("Warning: pytesseract not installed. OCR will not work.")

    def is_scanned_pdf(self, pdf_path: str) -> bool:
        try:
            doc = fitz.open(pdf_path)
            text_found = False
            for page in doc:
                if page.get_text().strip():
                    text_found = True
                    break
            doc.close()
            return not text_found
        except Exception:
            return False

    def convert(self, pdf_path: str, docx_path: str) -> tuple:
        try:
            is_scan = self.is_scanned_pdf(pdf_path)
            cv = Converter(pdf_path)
            
            if is_scan:
                mode = "PDF Scan (OCR Active)"
                cv.convert(docx_path, start=0, end=None, ocr=True)
            else:
                mode = "PDF Digital"
                cv.convert(docx_path, start=0, end=None)

            cv.close()
            
            # Post-process DOCX untuk format gambar dan caption
            self._post_process_docx(docx_path)
            
            return True, docx_path, mode
        except Exception as e:
            logger.error(f"Error converting: {e}")
            return False, str(e), "Unknown"

    # ══════════════════════════════════════════════════════════════════════
    # IMAGE PROCESSING - Convert all images to inline
    # ══════════════════════════════════════════════════════════════════════

    def _set_image_inline(self, run):
        """
        Mengubah gambar dari anchor (floating) menjadi inline with text.
        """
        drawing = run._element.find(qn('w:drawing'))
        if drawing is None:
            return
        
        anchor = drawing.find(qn('wp:anchor'))
        if anchor is None:
            return  # Sudah inline
        
        # Ambil elemen-elemen dari anchor
        extent = anchor.find(qn('wp:extent'))
        effect_extent = anchor.find(qn('wp:effectExtent'))
        docPr = anchor.find(qn('wp:docPr'))
        cNvGraphicFramePr = anchor.find(qn('wp:cNvGraphicFramePr'))
        graphic = anchor.find(qn('a:graphic'))
        
        if extent is None or docPr is None or graphic is None:
            return
        
        # Hapus anchor
        drawing.remove(anchor)
        
        # Buat elemen inline baru
        inline_elem = OxmlElement('wp:inline')
        inline_elem.set(qn('distT'), '0')
        inline_elem.set(qn('distB'), '0')
        inline_elem.set(qn('distL'), '0')
        inline_elem.set(qn('distR'), '0')
        
        # Copy extent
        new_extent = OxmlElement('wp:extent')
        new_extent.set(qn('cx'), extent.get(qn('cx'), '0'))
        new_extent.set(qn('cy'), extent.get(qn('cy'), '0'))
        inline_elem.append(new_extent)
        
        # Copy effectExtent jika ada
        if effect_extent is not None:
            new_effect = OxmlElement('wp:effectExtent')
            for attr in effect_extent.attrib:
                new_effect.set(attr, effect_extent.get(attr))
            inline_elem.append(new_effect)
        
        # Copy docPr
        new_docPr = copy.deepcopy(docPr)
        inline_elem.append(new_docPr)
        
        # Copy cNvGraphicFramePr jika ada
        if cNvGraphicFramePr is not None:
            new_cNv = copy.deepcopy(cNvGraphicFramePr)
            inline_elem.append(new_cNv)
        
        # Copy graphic
        new_graphic = copy.deepcopy(graphic)
        inline_elem.append(new_graphic)
        
        # Tambahkan inline ke drawing
        drawing.append(inline_elem)

    def _process_all_images_inline(self, doc):
        """
        Memproses semua gambar di dokumen menjadi inline.
        """
        # Proses paragraf utama
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                self._set_image_inline(run)
        
        # Proses gambar di dalam tabel
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            self._set_image_inline(run)

    # ══════════════════════════════════════════════════════════════════════
    # CAPTION DETECTION - Identifikasi paragraf caption gambar
    # ══════════════════════════════════════════════════════════════════════

    def _paragraph_has_image(self, paragraph):
        """Cek apakah paragraf mengandung gambar."""
        xml = paragraph._element.xml
        return 'wp:inline' in xml or 'wp:anchor' in xml

    def _is_figure_caption(self, text):
        """Cek apakah text adalah caption gambar (Figure X – ...)."""
        pattern = r'^Figure\s+\d+\s*[–\-]\s*.+'
        return bool(re.match(pattern, text.strip(), re.IGNORECASE))

    def _is_key_heading(self, text):
        """Cek apakah text adalah heading 'Key'."""
        return bool(re.match(r'^Key\s*$', text.strip(), re.IGNORECASE))

    def _is_key_item(self, text):
        """Cek apakah text adalah item dari Key (misal: '1 contact strip')."""
        pattern = r'^\d+\s+.+'
        return bool(re.match(pattern, text.strip()))

    def _is_note_paragraph(self, text):
        """Cek apakah text dimulai dengan NOTE."""
        pattern = r'^NOTE\s+\d*\s*.+'
        return bool(re.match(pattern, text.strip(), re.IGNORECASE))

    def _is_note_to_entry(self, text):
        """Cek apakah text adalah Note to entry."""
        pattern = r'^Note\s+\d+\s+to\s+entry\s*:.+'
        return bool(re.match(pattern, text.strip(), re.IGNORECASE))

    def _find_caption_blocks(self, doc):
        """
        Menemukan semua blok caption gambar beserta item-item terkait.
        Returns: list of (start_index, end_index, block_type)
        """
        captions = []
        paragraphs = doc.paragraphs
        n = len(paragraphs)
        i = 0
        
        while i < n:
            text = paragraphs[i].text.strip()
            
            # Lewati paragraf kosong atau yang mengandung gambar
            if not text or self._paragraph_has_image(paragraphs[i]):
                i += 1
                continue
            
            start_idx = None
            block_type = None
            
            # 1. Cek Figure caption
            if self._is_figure_caption(text):
                start_idx = i
                block_type = 'figure'
            
            # 2. Cek Key heading
            elif self._is_key_heading(text):
                start_idx = i
                block_type = 'key'
            
            # 3. Cek NOTE paragraph
            elif self._is_note_paragraph(text):
                start_idx = i
                block_type = 'note'
            
            # 4. Cek Note to entry
            elif self._is_note_to_entry(text):
                start_idx = i
                block_type = 'note_to_entry'
            
            if start_idx is not None:
                end_idx = start_idx
                
                # Jika figure caption, cek apakah ada Key setelahnya
                if block_type == 'figure':
                    j = start_idx + 1
                    # Lewati paragraf kosong
                    while j < n and not paragraphs[j].text.strip():
                        j += 1
                    
                    # Cek apakah paragraf berikutnya adalah Key
                    if j < n and self._is_key_heading(paragraphs[j].text.strip()):
                        # Temukan semua item Key
                        end_idx = j
                        k = j + 1
                        while k < n:
                            item_text = paragraphs[k].text.strip()
                            if self._is_key_item(item_text):
                                end_idx = k
                                k += 1
                            elif not item_text:  # Lewati baris kosong dalam blok Key
                                k += 1
                            else:
                                break
                
                # Jika Key heading, temukan semua item
                elif block_type == 'key':
                    j = start_idx + 1
                    while j < n:
                        item_text = paragraphs[j].text.strip()
                        if self._is_key_item(item_text):
                            end_idx = j
                            j += 1
                        elif not item_text:
                            j += 1
                        else:
                            break
                
                captions.append((start_idx, end_idx, block_type))
                i = end_idx + 1
            else:
                i += 1
        
        return captions

    # ══════════════════════════════════════════════════════════════════════
    # CAPTION FORMATTING - Format caption sesuai permintaan
    # ══════════════════════════════════════════════════════════════════════

    def _set_run_font(self, run, bold=False):
        """
        Set font untuk run: Arial, 10pt, bold/unbold.
        """
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.bold = bold
        
        # Set East Asia font juga
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')
        rFonts.set(qn('w:eastAsia'), 'Arial')
        rFonts.set(qn('w:cs'), 'Arial')

    def _set_paragraph_spacing(self, paragraph):
        """
        Set spacing paragraf: single, 0pt before/after.
        """
        pf = paragraph.paragraph_format
        pf.line_spacing = 1.0  # Single spacing
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        
        # Juga set via XML untuk memastikan
        pPr = paragraph._element.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), '240')  # Single spacing
        spacing.set(qn('w:lineRule'), 'auto')

    def _format_figure_caption(self, paragraph):
        """
        Format caption gambar: semua text unbold.
        """
        text = paragraph.text
        
        # Hapus semua run lama (kecuali yang mengandung gambar)
        runs_to_remove = []
        for run in paragraph.runs:
            if not self._paragraph_has_image(paragraph):
                runs_to_remove.append(run._element)
        
        for elem in runs_to_remove:
            paragraph._element.remove(elem)
        
        # Buat run baru dengan format yang benar
        if text:
            new_run = paragraph.add_run(text)
            self._set_run_font(new_run, bold=False)
        
        self._set_paragraph_spacing(paragraph)

    def _format_key_heading(self, paragraph):
        """
        Format heading 'Key': text "Key" bold.
        """
        text = paragraph.text.strip()
        
        # Hapus semua run lama
        runs_to_remove = [run._element for run in paragraph.runs]
        for elem in runs_to_remove:
            paragraph._element.remove(elem)
        
        # Buat run baru - "Key" bold
        if text:
            new_run = paragraph.add_run(text)
            self._set_run_font(new_run, bold=True)
        
        self._set_paragraph_spacing(paragraph)

    def _format_key_item(self, paragraph):
        """
        Format item Key (misal: '1 contact strip'): semua unbold.
        """
        text = paragraph.text
        
        # Hapus semua run lama
        runs_to_remove = [run._element for run in paragraph.runs]
        for elem in runs_to_remove:
            paragraph._element.remove(elem)
        
        # Buat run baru dengan format yang benar
        if text:
            new_run = paragraph.add_run(text)
            self._set_run_font(new_run, bold=False)
        
        self._set_paragraph_spacing(paragraph)

    def _format_note_paragraph(self, paragraph):
        """
        Format NOTE paragraph: 'NOTE' bold, sisanya unbold.
        """
        text = paragraph.text
        
        # Parse: "NOTE" atau "NOTE X" adalah bold, sisanya tidak
        match = re.match(r'^(NOTE\s*\d*)(.*)$', text, re.IGNORECASE | re.DOTALL)
        
        # Hapus semua run lama
        runs_to_remove = [run._element for run in paragraph.runs]
        for elem in runs_to_remove:
            paragraph._element.remove(elem)
        
        if match:
            note_word = match.group(1)  # "NOTE" atau "NOTE 1"
            rest_text = match.group(2)  # Sisa text
            
            # Run pertama: "NOTE" - BOLD
            if note_word:
                run1 = paragraph.add_run(note_word)
                self._set_run_font(run1, bold=True)
            
            # Run kedua: sisa text - UNBOLD
            if rest_text:
                run2 = paragraph.add_run(rest_text)
                self._set_run_font(run2, bold=False)
        
        self._set_paragraph_spacing(paragraph)

    def _format_note_to_entry(self, paragraph):
        """
        Format Note to entry: 'Note X to entry:' bold, sisanya unbold.
        """
        text = paragraph.text
        
        # Parse
        match = re.match(r'^(Note\s+\d+\s+to\s+entry\s*:)(.*)$', text, re.IGNORECASE | re.DOTALL)
        
        # Hapus semua run lama
        runs_to_remove = [run._element for run in paragraph.runs]
        for elem in runs_to_remove:
            paragraph._element.remove(elem)
        
        if match:
            note_header = match.group(1)
            rest_text = match.group(2)
            
            if note_header:
                run1 = paragraph.add_run(note_header)
                self._set_run_font(run1, bold=True)
            
            if rest_text:
                run2 = paragraph.add_run(rest_text)
                self._set_run_font(run2, bold=False)
        else:
            # Fallback: semua unbold
            if text:
                new_run = paragraph.add_run(text)
                self._set_run_font(new_run, bold=False)
        
        self._set_paragraph_spacing(paragraph)

    def _format_caption_block(self, paragraphs, start_idx, end_idx, block_type):
        """
        Format seluruh blok caption sesuai tipenya.
        """
        for i in range(start_idx, end_idx + 1):
            para = paragraphs[i]
            text = para.text.strip()
            
            if not text:
                continue
            
            if block_type == 'figure':
                # Cek apakah ini figure caption atau Key/item
                if self._is_figure_caption(text):
                    self._format_figure_caption(para)
                elif self._is_key_heading(text):
                    self._format_key_heading(para)
                elif self._is_key_item(text):
                    self._format_key_item(para)
            
            elif block_type == 'key':
                if self._is_key_heading(text):
                    self._format_key_heading(para)
                elif self._is_key_item(text):
                    self._format_key_item(para)
            
            elif block_type == 'note':
                self._format_note_paragraph(para)
            
            elif block_type == 'note_to_entry':
                self._format_note_to_entry(para)

    # ══════════════════════════════════════════════════════════════════════
    # MAIN POST-PROCESSING
    # ══════════════════════════════════════════════════════════════════════

    def _post_process_docx(self, docx_path: str):
        """
        Post-processing utama DOCX:
        1. Set semua gambar menjadi inline
        2. Format semua caption gambar
        """
        try:
            doc = Document(docx_path)
            
            # Step 1: Set semua gambar menjadi inline
            logger.info("Processing images to inline...")
            self._process_all_images_inline(doc)
            
            # Step 2: Identifikasi dan format semua caption
            logger.info("Processing figure captions...")
            caption_blocks = self._find_caption_blocks(doc)
            
            for start_idx, end_idx, block_type in caption_blocks:
                self._format_caption_block(doc.paragraphs, start_idx, end_idx, block_type)
            
            # Step 3: Cari dan format NOTE yang berdiri sendiri (bukan bagian figure)
            logger.info("Processing standalone NOTE paragraphs...")
            processed_indices = set()
            for start_idx, end_idx, _ in caption_blocks:
                for i in range(start_idx, end_idx + 1):
                    processed_indices.add(i)
            
            for i, para in enumerate(doc.paragraphs):
                if i in processed_indices:
                    continue
                text = para.text.strip()
                if self._is_note_paragraph(text):
                    self._format_note_paragraph(para)
                elif self._is_note_to_entry(text):
                    self._format_note_to_entry(para)
            
            # Save
            doc.save(docx_path)
            logger.info(f"Post-processing complete. {len(caption_blocks)} caption blocks processed.")
            
        except Exception as e:
            logger.error(f"Error in post-processing: {e}")
            raise
